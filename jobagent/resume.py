"""Resume ingestion: file -> text -> structured Profile."""
from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Optional

from .config import Config
from .models import Profile

log = logging.getLogger(__name__)

# Skills we can recognise without an LLM. Tilted towards DevOps/platform/SRE
# because that is what this agent is pointed at; extend freely in config.
KNOWN_SKILLS = [
    "kubernetes", "k8s", "docker", "terraform", "ansible", "puppet", "chef",
    "jenkins", "gitlab ci", "github actions", "circleci", "argocd", "argo cd",
    "flux", "helm", "istio", "linkerd", "envoy", "nginx", "haproxy",
    "aws", "azure", "gcp", "google cloud", "openstack", "vmware", "openshift",
    "ec2", "eks", "aks", "gke", "lambda", "s3", "rds", "cloudformation", "cdk",
    "pulumi", "packer", "vault", "consul", "nomad", "hashicorp",
    "prometheus", "grafana", "datadog", "new relic", "splunk", "elk",
    "elasticsearch", "logstash", "kibana", "opentelemetry", "jaeger", "loki",
    "pagerduty", "sentry", "cloudwatch", "zabbix", "nagios",
    "python", "go", "golang", "bash", "shell", "ruby", "java", "javascript",
    "typescript", "groovy", "powershell", "yaml", "json", "sql",
    "linux", "ubuntu", "centos", "rhel", "debian", "networking", "tcp/ip",
    "dns", "load balancing", "cdn", "cloudflare", "vpn", "firewall",
    "ci/cd", "cicd", "devops", "sre", "gitops", "iac", "infrastructure as code",
    "microservices", "serverless", "observability", "monitoring", "logging",
    "incident response", "on-call", "sla", "slo", "sli", "chaos engineering",
    "postgresql", "mysql", "mongodb", "redis", "cassandra", "kafka",
    "rabbitmq", "dynamodb", "snowflake", "airflow", "spark",
    "security", "devsecops", "compliance", "soc2", "iso 27001", "pci",
    "cost optimization", "finops", "capacity planning", "disaster recovery",
    "high availability", "scalability", "automation", "platform engineering",
    "agile", "scrum", "mentoring", "team leadership", "stakeholder management",
]

CERT_PATTERNS = [
    r"AWS Certified [A-Za-z \-]+", r"Certified Kubernetes [A-Za-z ]+",
    r"\bCKA\b", r"\bCKAD\b", r"\bCKS\b", r"Azure (?:Solutions )?Architect[A-Za-z ]*",
    r"Google Cloud (?:Professional |Certified )?[A-Za-z ]+",
    r"HashiCorp Certified[A-Za-z \-]*", r"\bRHCE\b", r"\bRHCSA\b", r"\bITIL\b",
    r"Terraform Associate", r"\bPMP\b",
]

_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
# Deliberately loose; the 9-15 digit check below is what actually validates it.
_PHONE = re.compile(r"(?:\+\d{1,3}[\s.\-]?)?(?:\(\d{2,5}\)[\s.\-]?)?\d[\d\s.\-]{5,14}\d")
_LINKEDIN = re.compile(r"(?:https?://)?(?:[\w]{2,3}\.)?linkedin\.com/in/[\w\-%]+", re.I)
_GITHUB = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-]+", re.I)


def extract_text(path: str | Path) -> str:
    """Pull plain text out of a PDF, DOCX, TXT or MD resume."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Resume not found: {path}")
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF resume needs pypdf: pip install pypdf") from exc
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()

    if suffix in {".docx", ".doc"}:
        try:
            import docx
        except ImportError as exc:
            raise RuntimeError("DOCX resume needs python-docx: pip install python-docx") from exc
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(p for p in parts if p.strip()).strip()

    if suffix in {".txt", ".md", ".markdown", ""}:
        return path.read_text(encoding="utf-8", errors="replace").strip()

    raise ValueError(f"Unsupported resume format: {suffix} (use pdf, docx, txt or md)")


def parse_heuristic(text: str, extra_skills: Optional[list[str]] = None) -> Profile:
    """Regex/keyword extraction. Always runs; the LLM only refines it."""
    profile = Profile(resume_text=text)
    lowered = text.lower()

    if m := _EMAIL.search(text):
        profile.email = m.group(0)
    if m := _LINKEDIN.search(text):
        profile.linkedin_url = _https(m.group(0))
    if m := _GITHUB.search(text):
        profile.github_url = _https(m.group(0))

    # Phone: scan the header, where it almost always sits, to avoid matching dates.
    header = "\n".join(text.splitlines()[:12])
    for candidate in _PHONE.findall(header):
        digits = re.sub(r"\D", "", candidate)
        if 9 <= len(digits) <= 15:
            profile.phone = candidate.strip()
            break

    # Name: first non-empty line that isn't a contact detail or a section header.
    for line in text.splitlines()[:8]:
        line = line.strip()
        if not line or "@" in line or _PHONE.fullmatch(line) or len(line) > 60:
            continue
        if re.search(r"(?i)\b(resume|curriculum vitae|cv)\b", line):
            continue
        words = line.split()
        if 1 < len(words) <= 5 and not any(ch.isdigit() for ch in line):
            profile.full_name = line
            break

    vocabulary = list(KNOWN_SKILLS) + [s.lower() for s in (extra_skills or [])]
    found = {s for s in vocabulary if re.search(rf"(?<![\w+#]){re.escape(s)}(?![\w+#])", lowered)}
    profile.skills = sorted(found)

    certs: list[str] = []
    for pattern in CERT_PATTERNS:
        certs.extend(m.strip() for m in re.findall(pattern, text))
    profile.certifications = sorted(set(certs))

    if m := re.search(r"(\d{1,2})\+?\s*years?(?:\s+of)?\s+(?:of\s+)?experience", lowered):
        profile.years_experience = float(m.group(1))
    else:
        profile.years_experience = _years_from_date_ranges(text)

    profile.titles = _find_titles(text)
    return profile


def _https(url: str) -> str:
    return url if url.startswith("http") else f"https://{url}"


def _years_from_date_ranges(text: str) -> Optional[float]:
    """Infer tenure from the earliest 4-digit year that looks like a job start."""
    from datetime import datetime
    years = [int(y) for y in re.findall(r"\b(19[89]\d|20[0-4]\d)\b", text)]
    plausible = [y for y in years if 1985 <= y <= datetime.now().year]
    if not plausible:
        return None
    span = datetime.now().year - min(plausible)
    return float(span) if 0 < span < 45 else None


def _find_titles(text: str) -> list[str]:
    pattern = re.compile(
        r"(?im)^\s*((?:lead|principal|staff|senior|sr\.?|head of|director of|"
        r"associate)?\s*(?:devops|sre|site reliability|platform|cloud|"
        r"infrastructure|systems?|software|security)\s*"
        r"(?:engineer|architect|manager|lead|consultant|administrator|specialist))\b"
    )
    seen: list[str] = []
    for match in pattern.findall(text):
        title = re.sub(r"\s+", " ", match).strip().title()
        if title and title not in seen:
            seen.append(title)
    return seen[:12]


def build_profile(config: Config, llm=None) -> Profile:
    """Full pipeline: read the resume, parse it, refine with the LLM, apply overrides.

    Config always wins over anything extracted - it is the thing you typed by
    hand, and a wrong phone number on 200 applications is expensive.
    """
    conf = config.profile or {}
    resume_path = conf.get("resume_path", "")
    text = ""
    if resume_path and Path(resume_path).exists():
        text = extract_text(resume_path)
    elif resume_path:
        log.warning("resume_path %s does not exist", resume_path)

    profile = parse_heuristic(text, conf.get("extra_skills"))
    profile.resume_path = resume_path

    if text and llm is not None and config.match.use_llm:
        try:
            refine_with_llm(profile, llm)
        except Exception as exc:  # noqa: BLE001 - heuristics already gave us a profile
            log.warning("LLM resume parse failed, keeping heuristic profile: %s", exc)

    for field in ("full_name", "email", "phone", "location", "headline", "summary",
                  "linkedin_url", "github_url", "portfolio_url"):
        if conf.get(field):
            setattr(profile, field, conf[field])
    if conf.get("years_experience"):
        profile.years_experience = float(conf["years_experience"])
    if conf.get("extra_skills"):
        profile.skills = sorted(set(profile.skills) | {s.lower() for s in conf["extra_skills"]})
    profile.answers = dict(conf.get("answers") or {})
    return profile


RESUME_PARSE_PROMPT = """Extract structured data from this resume. Return ONLY a JSON object, no prose.

Keys:
- full_name, email, phone, location (city, country), headline (current role in <=12 words)
- summary (2-3 sentences, written in third person, factual)
- years_experience (number)
- skills (array of lowercase technology/tool names actually used, max 40)
- titles (array of past job titles, most recent first)
- companies (array of employer names, most recent first)
- education (array of "Degree, Institution, Year" strings)
- certifications (array)

Use "" or [] for anything not present. Never invent facts.

RESUME:
{text}
"""


def refine_with_llm(profile: Profile, llm) -> Profile:
    """Overlay LLM-extracted fields onto the heuristic profile."""
    raw = llm.complete(
        RESUME_PARSE_PROMPT.format(text=profile.resume_text[:18000]),
        max_tokens=2000,
        system="You extract structured data from resumes and return strict JSON.",
    )
    data = _json_from(raw)
    if not data:
        return profile

    for key in ("full_name", "email", "phone", "location", "headline", "summary"):
        if data.get(key):
            setattr(profile, key, str(data[key]).strip())
    if data.get("years_experience"):
        try:
            profile.years_experience = float(data["years_experience"])
        except (TypeError, ValueError):
            pass
    if data.get("skills"):
        merged = {s.lower().strip() for s in data["skills"] if isinstance(s, str)}
        profile.skills = sorted(set(profile.skills) | merged)
    for key in ("titles", "companies", "education", "certifications"):
        values = [str(v).strip() for v in (data.get(key) or []) if str(v).strip()]
        if values:
            setattr(profile, key, values)
    return profile


def _json_from(raw: str) -> dict:
    """Pull a JSON object out of a model response that may be fenced or chatty."""
    if not raw:
        return {}
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.S)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    start, depth = text.find("{"), 0
    if start == -1:
        return {}
    for i, ch in enumerate(text[start:], start):
        depth += (ch == "{") - (ch == "}")
        if depth == 0:
            try:
                return json.loads(text[start:i + 1])
            except json.JSONDecodeError:
                return {}
    return {}
